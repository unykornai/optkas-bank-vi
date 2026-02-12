"""
Output Exporter
===============
Converts rendered Markdown agreements to DOCX and PDF formats
using python-docx and pypandoc.
"""

from __future__ import annotations

import subprocess
from pathlib import Path

from engine.schema_loader import ensure_output_dir


def export_markdown(content: str, filename: str) -> Path:
    """Write rendered Markdown to the output directory."""
    out_dir = ensure_output_dir()
    path = out_dir / f"{filename}.md"
    path.write_text(content, encoding="utf-8")
    return path


def export_docx(markdown_path: Path, output_name: str | None = None) -> Path:
    """
    Convert Markdown to DOCX using pypandoc.

    Falls back to python-docx if pandoc is not available.
    """
    out_dir = ensure_output_dir()
    name = output_name or markdown_path.stem
    docx_path = out_dir / f"{name}.docx"

    try:
        import pypandoc

        pypandoc.convert_file(
            str(markdown_path),
            "docx",
            outputfile=str(docx_path),
            extra_args=[
                "--standalone",
                "--reference-doc=" + str(_get_reference_docx()),
            ] if _get_reference_docx().exists() else ["--standalone"],
        )
    except (ImportError, OSError):
        # Fallback: use python-docx for basic conversion
        _markdown_to_docx_fallback(markdown_path, docx_path)

    return docx_path


def export_pdf(markdown_path: Path, output_name: str | None = None) -> Path:
    """
    Convert Markdown to PDF using pypandoc/pandoc.

    Requires pandoc and a PDF engine (e.g., pdflatex, wkhtmltopdf).
    """
    out_dir = ensure_output_dir()
    name = output_name or markdown_path.stem
    pdf_path = out_dir / f"{name}.pdf"

    try:
        import pypandoc

        pypandoc.convert_file(
            str(markdown_path),
            "pdf",
            outputfile=str(pdf_path),
            extra_args=[
                "--pdf-engine=xelatex",
                "-V", "geometry:margin=1in",
                "-V", "fontsize=11pt",
            ],
        )
    except (ImportError, OSError) as e:
        raise RuntimeError(
            f"PDF export requires pandoc and a PDF engine (xelatex or wkhtmltopdf). "
            f"Install with: choco install pandoc miktex  |  Error: {e}"
        )

    return pdf_path


def _get_reference_docx() -> Path:
    """Get path to reference DOCX template for styling."""
    return Path(__file__).parent.parent / "templates" / "reference.docx"


def _markdown_to_docx_fallback(md_path: Path, docx_path: Path) -> None:
    """Basic Markdown-to-DOCX conversion using python-docx."""
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(11)

    content = md_path.read_text(encoding="utf-8")

    for line in content.split("\n"):
        stripped = line.strip()

        if stripped.startswith("# "):
            p = doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith("## "):
            p = doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("### "):
            p = doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith("---"):
            doc.add_page_break()
        elif stripped.startswith("- "):
            doc.add_paragraph(stripped[2:], style="List Bullet")
        elif stripped.startswith("("):
            doc.add_paragraph(stripped, style="List Number")
        elif stripped.startswith("**") and stripped.endswith("**"):
            p = doc.add_paragraph()
            run = p.add_run(stripped.strip("*"))
            run.bold = True
        elif stripped == "":
            doc.add_paragraph("")
        else:
            doc.add_paragraph(stripped)

    doc.save(str(docx_path))
